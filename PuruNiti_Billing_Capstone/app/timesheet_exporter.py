import os
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_timesheet_pdf(ts, output_path):
    """Generates a professional PDF timesheet ledger using ReportLab."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TSTitle', parent=styles['Heading1'],
        fontName='Helvetica-Bold', fontSize=18, leading=22,
        textColor=colors.HexColor('#1A365D'), spaceAfter=15
    )
    section_title = ParagraphStyle(
        'TSSection', parent=styles['Heading2'],
        fontName='Helvetica-Bold', fontSize=11, leading=14,
        textColor=colors.HexColor('#2C3E50'), spaceAfter=8
    )
    body_style = ParagraphStyle(
        'TSBody', parent=styles['Normal'],
        fontName='Helvetica', fontSize=9, leading=12,
        textColor=colors.HexColor('#2D3748')
    )
    header_style = ParagraphStyle(
        'TSHeader', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=9.5, leading=12,
        textColor=colors.white
    )
    
    story = []
    
    # Title
    story.append(Paragraph("TIMESHEET & BILLING SUMMARY", title_style))
    story.append(Spacer(1, 10))
    
    # Metadata block
    meta_data = [
        [Paragraph(f"<b>Client:</b> {ts.get('client_name')}", body_style), 
         Paragraph(f"<b>Timesheet ID:</b> {ts.get('timesheet_id')}", body_style)],
        [Paragraph(f"<b>Period:</b> {ts.get('start_date')} to {ts.get('end_date')}", body_style), 
         Paragraph(f"<b>Status:</b> {ts.get('status')}", body_style)],
        [Paragraph(f"<b>Total Hours:</b> {ts.get('total_hours'):.2f} hrs", body_style), 
         Paragraph(f"<b>Total Amount:</b> ₹ {ts.get('total_amount'):,.2f}", body_style)]
    ]
    
    meta_table = Table(meta_data, colWidths=[240, 240])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 20))
    
    # Entries Header
    story.append(Paragraph("Activity Log Entries", section_title))
    
    # Table headers & data
    table_data = [[
        Paragraph("Date", header_style),
        Paragraph("Activity / Task Description", header_style),
        Paragraph("Hours", header_style),
        Paragraph("Rate (₹/hr)", header_style),
        Paragraph("Line Total (₹)", header_style)
    ]]
    
    for entry in ts.get("entries", []):
        table_data.append([
            Paragraph(entry.get("date", ""), body_style),
            Paragraph(entry.get("activity", ""), body_style),
            Paragraph(f"{entry.get('hours', 0.0):.2f}", body_style),
            Paragraph(f"{entry.get('rate', 0.0):.2f}", body_style),
            Paragraph(f"{entry.get('line_total', 0.0):.2f}", body_style)
        ])
        
    # Append Total Row
    table_data.append([
        Paragraph("<b>Total Summary</b>", body_style),
        Paragraph("", body_style),
        Paragraph(f"<b>{ts.get('total_hours'):.2f}</b>", body_style),
        Paragraph("", body_style),
        Paragraph(f"<b>₹ {ts.get('total_amount'):,.2f}</b>", body_style)
    ])
    
    entries_table = Table(table_data, colWidths=[70, 210, 60, 60, 80])
    entries_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1A365D')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('SPAN', (0, -1), (1, -1)), # Span total label
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#F8FAFC')),
    ]))
    story.append(entries_table)
    story.append(Spacer(1, 40))
    
    # Signatures
    sig_data = [
        [Paragraph("<b>Prepared By:</b>", body_style), Paragraph("<b>Approved By (Client):</b>", body_style)],
        [Spacer(1, 30), Spacer(1, 30)],
        [Paragraph("_____________________________<br/>Consultant Signature", body_style), 
         Paragraph("_____________________________<br/>Authorized Signatory", body_style)]
    ]
    sig_table = Table(sig_data, colWidths=[240, 240])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(sig_table)
    
    doc.build(story)
    return True

def export_timesheet_excel(ts, output_path):
    """Generates a professional Excel timesheet template using openpyxl."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Timesheet Log"
    
    # Title Block
    ws.cell(row=1, column=1, value="TIMESHEET & BILLING SUMMARY").font = Font(name="Segoe UI", size=14, bold=True, color="1A365D")
    
    # Metadata
    meta_items = [
        ("Client Name:", ts.get("client_name"), "Timesheet ID:", ts.get("timesheet_id")),
        ("Period:", f"{ts.get('start_date')} to {ts.get('end_date')}", "Status:", ts.get("status")),
        ("Total Hours:", ts.get("total_hours"), "Total Amount:", ts.get("total_amount"))
    ]
    
    for idx, row_vals in enumerate(meta_items):
        r = 3 + idx
        ws.cell(row=r, column=1, value=row_vals[0]).font = Font(name="Segoe UI", size=9.5, bold=True)
        ws.cell(row=r, column=2, value=row_vals[1]).font = Font(name="Segoe UI", size=9.5)
        ws.cell(row=r, column=4, value=row_vals[2]).font = Font(name="Segoe UI", size=9.5, bold=True)
        ws.cell(row=r, column=5, value=row_vals[3]).font = Font(name="Segoe UI", size=9.5)
        if idx == 2:
            ws.cell(row=r, column=5).number_format = '[$₹-3601] #,##0.00'
            
    # Table Headers
    headers = ["Date", "Activity / Task Description", "Hours", "Hourly Rate (₹/hr)", "Line Total (₹)"]
    ws.append([]) # spacer
    ws.append([]) # spacer
    ws.append(headers)
    
    h_row = 7
    header_fill = PatternFill(start_color="1A365D", end_color="1A365D", fill_type="solid")
    header_font = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
    thin = Side(border_style="thin", color="CBD5E1")
    grid_border = Border(top=thin, left=thin, right=thin, bottom=thin)
    
    for col_idx in range(1, 6):
        cell = ws.cell(row=h_row, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center" if col_idx != 2 else "left", vertical="center")
        cell.border = grid_border
        
    ws.row_dimensions[h_row].height = 24
    
    # Entries Rows
    start_r = 8
    for entry in ts.get("entries", []):
        row_data = [
            entry.get("date", ""),
            entry.get("activity", ""),
            float(entry.get("hours", 0.0)),
            float(entry.get("rate", 0.0)),
            float(entry.get("line_total", 0.0))
        ]
        ws.append(row_data)
        curr_r = ws.max_row
        
        ws.row_dimensions[curr_r].height = 18
        for col_idx in range(1, 6):
            c = ws.cell(row=curr_r, column=col_idx)
            c.font = Font(name="Segoe UI", size=9.5)
            c.border = grid_border
            if col_idx == 1:
                c.alignment = Alignment(horizontal="center")
            elif col_idx == 3:
                c.alignment = Alignment(horizontal="right")
                c.number_format = '0.00'
            elif col_idx in [4, 5]:
                c.alignment = Alignment(horizontal="right")
                c.number_format = '#,##0.00'
                
    # Total Summary Row
    end_r = ws.max_row
    tot_row = end_r + 1
    ws.cell(row=tot_row, column=1, value="Total Summary").font = Font(name="Segoe UI", size=9.5, bold=True)
    ws.cell(row=tot_row, column=3, value=f"=SUM(C8:C{end_r})").font = Font(name="Segoe UI", size=9.5, bold=True)
    ws.cell(row=tot_row, column=3).number_format = '0.00'
    ws.cell(row=tot_row, column=3).alignment = Alignment(horizontal="right")
    
    ws.cell(row=tot_row, column=5, value=f"=SUM(E8:E{end_r})").font = Font(name="Segoe UI", size=9.5, bold=True)
    ws.cell(row=tot_row, column=5).number_format = '[$₹-3601] #,##0.00'
    ws.cell(row=tot_row, column=5).alignment = Alignment(horizontal="right")
    
    total_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    for col_idx in range(1, 6):
        c = ws.cell(row=tot_row, column=col_idx)
        c.fill = total_fill
        c.border = grid_border
        
    ws.row_dimensions[tot_row].height = 20
    
    # Autofit columns
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)
        
    ws.column_dimensions['B'].width = 35
    
    wb.save(output_path)
    return True

def export_timesheet_word(ts, output_path):
    """Generates an editable Microsoft Word timesheet ledger using python-docx."""
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Document Title
    title = doc.add_paragraph()
    title_run = title.add_run("TIMESHEET & BILLING SUMMARY")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = docx.shared.RGBColor(26, 54, 93)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Metadata
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Client Name: ").bold = True
    meta_p.add_run(f"{ts.get('client_name')}\n")
    meta_p.add_run(f"Timesheet ID: ").bold = True
    meta_p.add_run(f"{ts.get('timesheet_id')}\n")
    meta_p.add_run(f"Period: ").bold = True
    meta_p.add_run(f"{ts.get('start_date')} to {ts.get('end_date')}\n")
    meta_p.add_run(f"Total Billing Hours: ").bold = True
    meta_p.add_run(f"{ts.get('total_hours'):.2f} hrs\n")
    meta_p.add_run(f"Total Amount: ").bold = True
    meta_p.add_run(f"₹ {ts.get('total_amount'):,.2f}\n")
    meta_p.add_run(f"Status: ").bold = True
    meta_p.add_run(f"{ts.get('status')}")
    
    doc.add_paragraph().add_run("Activity Log Entries").bold = True
    
    # Create Table (Headers + Data + Total)
    entries = ts.get("entries", [])
    table = doc.add_table(rows=len(entries) + 2, cols=5)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ["Date", "Activity Description", "Hours", "Rate (₹/hr)", "Total (₹)"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    for r_idx, entry in enumerate(entries):
        row_cells = table.rows[r_idx + 1].cells
        row_cells[0].text = entry.get("date", "")
        row_cells[1].text = entry.get("activity", "")
        row_cells[2].text = f"{entry.get('hours', 0.0):.2f}"
        row_cells[3].text = f"{entry.get('rate', 0.0):.2f}"
        row_cells[4].text = f"{entry.get('line_total', 0.0):.2f}"
        
    # Totals row
    tot_cells = table.rows[-1].cells
    tot_cells[0].text = "Total Summary"
    tot_cells[0].paragraphs[0].runs[0].font.bold = True
    tot_cells[2].text = f"{ts.get('total_hours'):.2f}"
    tot_cells[2].paragraphs[0].runs[0].font.bold = True
    tot_cells[4].text = f"₹ {ts.get('total_amount'):,.2f}"
    tot_cells[4].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph() # spacer
    doc.add_paragraph() # spacer
    
    # Signatures
    sig_p = doc.add_paragraph()
    sig_p.add_run("Prepared By:                                                     Approved By (Client):\n\n\n\n").bold = True
    sig_p.add_run("_____________________________                    _____________________________\n").bold = True
    sig_p.add_run("Consultant Signature                                          Authorized Signatory")
    
    doc.save(output_path)
    return True
