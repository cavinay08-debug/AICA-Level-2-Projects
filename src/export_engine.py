import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import os
import sqlite3
from src.statement_generator import generate_financial_statements, get_company_profile, generate_notes
from src.reconciliation_engine import run_reconciliation_check
from src.mapping_engine import get_mapping_register
from src.db import get_connection

def export_excel_working_papers(output_path):
    data = generate_financial_statements()
    notes_data = generate_notes()
    profile = data['profile']
    unit = data['unit']
    cy_label = profile.get('financial_year', '2024')
    py_label = profile.get('comparative_year', '2023')
    
    # Extract clean year strings (e.g. 2024, 2023 or 2015, 2014)
    cy_year = cy_label.split('-')[1] if '-' in cy_label and len(cy_label.split('-')[1]) == 4 else (f"20{cy_label.split('-')[1]}" if '-' in cy_label else cy_label)
    py_year = py_label.split('-')[1] if '-' in py_label and len(py_label.split('-')[1]) == 4 else (f"20{py_label.split('-')[1]}" if '-' in py_label else py_label)

    wb = openpyxl.Workbook()

    # ═══════════════════════════════════════════════════════════════════════
    # CORPORATE ICAI ANNUAL REPORT STYLING DEFINITIONS (Matched to Sample)
    # ═══════════════════════════════════════════════════════════════════════
    COLOR_PRIMARY_BLUE = "004C97"      # Dark Royal Blue (Main Headings / Key totals)
    COLOR_CY_FILL      = "DCE6F1"      # Soft Ice-Blue column tint for Current Year
    COLOR_TEXT_MAIN    = "1E293B"      # Dark Slate/Charcoal text
    COLOR_TEXT_MUTED   = "64748B"      # Subdued grey
    COLOR_BORDER_THIN  = "000000"      # Standard thin border
    COLOR_LINE_LIGHT   = "CBD5E1"      # Light grey separator

    font_title         = Font(name="Calibri", size=18, bold=True, color=COLOR_PRIMARY_BLUE)
    font_subtitle      = Font(name="Calibri", size=12, color="334155")
    font_unit_header   = Font(name="Calibri", size=10, italic=True, bold=True, color=COLOR_PRIMARY_BLUE)
    
    font_tbl_hdr_cy    = Font(name="Calibri", size=10, bold=True, color=COLOR_PRIMARY_BLUE)
    font_tbl_hdr_py    = Font(name="Calibri", size=10, bold=True, color=COLOR_TEXT_MAIN)
    font_tbl_hdr_note  = Font(name="Calibri", size=10, bold=True, color=COLOR_TEXT_MAIN)
    
    font_major_sec     = Font(name="Calibri", size=10.5, bold=True, color=COLOR_PRIMARY_BLUE)
    font_sub_head      = Font(name="Calibri", size=10.5, bold=True, color=COLOR_PRIMARY_BLUE)
    font_item_reg      = Font(name="Calibri", size=10.5, color=COLOR_TEXT_MAIN)
    font_item_sub      = Font(name="Calibri", size=10.0, italic=True, color=COLOR_TEXT_MAIN)
    
    font_cy_data       = Font(name="Calibri", size=10.5, bold=True, color=COLOR_PRIMARY_BLUE)
    font_py_data       = Font(name="Calibri", size=10.5, color=COLOR_TEXT_MAIN)
    
    font_total_lbl     = Font(name="Calibri", size=11, bold=True, color=COLOR_PRIMARY_BLUE)
    font_total_cy      = Font(name="Calibri", size=11, bold=True, color=COLOR_PRIMARY_BLUE)
    font_total_py      = Font(name="Calibri", size=11, bold=True, color="000000")
    
    font_sig_title     = Font(name="Calibri", size=10, bold=True, color=COLOR_TEXT_MAIN)
    font_sig_role      = Font(name="Calibri", size=9.5, color=COLOR_TEXT_MUTED)

    fill_cy_col        = PatternFill(start_color=COLOR_CY_FILL, end_color=COLOR_CY_FILL, fill_type="solid")
    fill_soft_note     = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")

    border_top_medium  = Border(top=Side(style='medium', color=COLOR_BORDER_THIN))
    border_bot_thin    = Border(bottom=Side(style='thin', color=COLOR_BORDER_THIN))
    border_tbl_hdr     = Border(top=Side(style='medium', color=COLOR_BORDER_THIN), bottom=Side(style='thin', color=COLOR_BORDER_THIN))
    border_subtotal    = Border(top=Side(style='thin', color=COLOR_BORDER_THIN), bottom=Side(style='thin', color=COLOR_BORDER_THIN))
    border_total_dbl   = Border(top=Side(style='thin', color=COLOR_BORDER_THIN), bottom=Side(style='double', color=COLOR_BORDER_THIN))

    NUM_FMT = '#,##0.00;(#,##0.00);"-"'

    def _add_statutory_signatures(ws, max_cols=4):
        ws.append([])
        ws.append(["Significant accounting policies", "1", "", ""])
        r_pol = ws.max_row
        ws.cell(r_pol, 1).font = font_major_sec
        ws.cell(r_pol, 2).font = font_tbl_hdr_note
        ws.cell(r_pol, 2).alignment = Alignment(horizontal="center")
        
        ws.append(["Contingent liabilities, capital and other commitments", "23", "", ""])
        r_con = ws.max_row
        ws.cell(r_con, 1).font = font_major_sec
        ws.cell(r_con, 2).font = font_tbl_hdr_note
        ws.cell(r_con, 2).alignment = Alignment(horizontal="center")
        
        ws.append([])
        ws.append(["The accompanying notes are an integral part of these financial statements", "", "", ""])
        ws.cell(ws.max_row, 1).font = Font(name="Calibri", size=10, bold=True, italic=True, color=COLOR_TEXT_MAIN)
        
        # Border divider
        r_div = ws.max_row
        for c in range(1, max_cols + 1):
            ws.cell(r_div, c).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))
            
        ws.append([])
        ws.append(["As per our report of even date", "", "For and on behalf of Board of Directors", ""])
        r_sig_h = ws.max_row
        ws.cell(r_sig_h, 1).font = font_sig_title
        ws.cell(r_sig_h, 3).font = font_sig_title
        
        ws.append(["For B S R & Co. LLP", "", "Managing Director and CEO", "Executive Director & CFO"])
        r_sig1 = ws.max_row
        ws.cell(r_sig1, 1).font = font_sig_title
        ws.cell(r_sig1, 3).font = font_sig_title
        ws.cell(r_sig1, 4).font = font_sig_title
        
        ws.append(["Chartered Accountants", "", "[DIN: 06699923]", "[DIN: 02762983]"])
        r_sig2 = ws.max_row
        ws.cell(r_sig2, 1).font = font_sig_role
        ws.cell(r_sig2, 3).font = font_sig_role
        ws.cell(r_sig2, 4).font = font_sig_role
        
        ws.append(["Firm Registration No. 101248W/W - 100022", "", "", ""])
        ws.cell(ws.max_row, 1).font = font_sig_role
        
        ws.append(["Partner", "", "Chairman - Audit Committee", "Company Secretary"])
        r_sig3 = ws.max_row
        ws.cell(r_sig3, 1).font = font_sig_title
        ws.cell(r_sig3, 3).font = font_sig_title
        ws.cell(r_sig3, 4).font = font_sig_title

    # ═══════════════════════════════════════════════════════════════════════
    # 1. BALANCE SHEET SHEET (Matches Sample Image 1)
    # ═══════════════════════════════════════════════════════════════════════
    ws_bs = wb.active
    ws_bs.title = "Balance Sheet"
    
    # Title Block
    ws_bs.append(["BALANCE SHEET"])
    ws_bs.cell(1, 1).font = font_title
    
    ws_bs.append([f"As at 31st March, {cy_year}"])
    ws_bs.cell(2, 1).font = font_subtitle
    
    # Blue underline below title
    ws_bs.cell(2, 1).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))
    ws_bs.cell(2, 2).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))
    ws_bs.cell(2, 3).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))
    ws_bs.cell(2, 4).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))

    # Unit note (right aligned)
    ws_bs.append([])
    ws_bs.append(["", "", "", f"(All amounts in {profile['currency']} {unit}, unless otherwise stated)"])
    ws_bs.cell(4, 4).font = font_unit_header
    ws_bs.cell(4, 4).alignment = Alignment(horizontal="right")

    # Table Header Row 5 & 6
    ws_bs.append(["", "Note", f"As at", f"As at"])
    ws_bs.append(["", "", f"31st March, {cy_year}", f"31st March, {py_year}"])
    
    # Format Table Header
    for r in (5, 6):
        ws_bs.cell(r, 2).font = font_tbl_hdr_note
        ws_bs.cell(r, 2).alignment = Alignment(horizontal="center", vertical="center")
        ws_bs.cell(r, 3).font = font_tbl_hdr_cy
        ws_bs.cell(r, 3).fill = fill_cy_col
        ws_bs.cell(r, 3).alignment = Alignment(horizontal="right", vertical="center")
        ws_bs.cell(r, 4).font = font_tbl_hdr_py
        ws_bs.cell(r, 4).alignment = Alignment(horizontal="right", vertical="center")

    for c in range(1, 5):
        ws_bs.cell(5, c).border = Border(top=Side(style='medium', color=COLOR_BORDER_THIN))
        ws_bs.cell(6, c).border = Border(bottom=Side(style='thin', color=COLOR_BORDER_THIN))

    # Data Rows
    bs_items = (
        data['bs_equity_liabilities'] +
        [{'title': 'TOTAL', 'is_total': True, 'cy': data['total_eq_liab_cy'], 'py': data['total_eq_liab_py']}] +
        data['bs_assets'] +
        [{'title': 'TOTAL', 'is_total': True, 'cy': data['total_assets_cy'], 'py': data['total_assets_py']}]
    )

    for item in bs_items:
        title = item.get('title', '')
        note = item.get('note', '')
        cy_v = item.get('cy', '')
        py_v = item.get('py', '')

        if item.get('is_header'):
            ws_bs.append([title, "", "", ""])
            r = ws_bs.max_row
            ws_bs.cell(r, 1).font = font_major_sec
            ws_bs.cell(r, 3).fill = fill_cy_col
        elif item.get('is_subheader'):
            ws_bs.append([f"  {title}", "", "", ""])
            r = ws_bs.max_row
            ws_bs.cell(r, 1).font = font_sub_head
            ws_bs.cell(r, 3).fill = fill_cy_col
        elif item.get('is_sublabel'):
            ws_bs.append([f"    {title}", "", "", ""])
            r = ws_bs.max_row
            ws_bs.cell(r, 1).font = font_item_sub
            ws_bs.cell(r, 3).fill = fill_cy_col
        elif item.get('is_total'):
            ws_bs.append([title, "", cy_v, py_v])
            r = ws_bs.max_row
            ws_bs.cell(r, 1).font = font_total_lbl
            ws_bs.cell(r, 3).font = font_total_cy
            ws_bs.cell(r, 3).fill = fill_cy_col
            ws_bs.cell(r, 3).number_format = NUM_FMT
            ws_bs.cell(r, 3).alignment = Alignment(horizontal="right")
            ws_bs.cell(r, 4).font = font_total_py
            ws_bs.cell(r, 4).number_format = NUM_FMT
            ws_bs.cell(r, 4).alignment = Alignment(horizontal="right")
            for c in range(1, 5):
                ws_bs.cell(r, c).border = border_total_dbl
            ws_bs.append([]) # space after total
        else:
            indent_prefix = "      " if not title.startswith(" ") else "    "
            ws_bs.append([f"{indent_prefix}{title.strip()}", note, cy_v, py_v])
            r = ws_bs.max_row
            ws_bs.cell(r, 1).font = font_item_reg
            ws_bs.cell(r, 2).font = font_tbl_hdr_note
            ws_bs.cell(r, 2).alignment = Alignment(horizontal="center")
            ws_bs.cell(r, 3).font = font_cy_data
            ws_bs.cell(r, 3).fill = fill_cy_col
            ws_bs.cell(r, 3).number_format = NUM_FMT
            ws_bs.cell(r, 3).alignment = Alignment(horizontal="right")
            ws_bs.cell(r, 4).font = font_py_data
            ws_bs.cell(r, 4).number_format = NUM_FMT
            ws_bs.cell(r, 4).alignment = Alignment(horizontal="right")

    _add_statutory_signatures(ws_bs, max_cols=4)

    # ═══════════════════════════════════════════════════════════════════════
    # 2. STATEMENT OF PROFIT AND LOSS (Matches Sample Image 2)
    # ═══════════════════════════════════════════════════════════════════════
    ws_pl = wb.create_sheet(title="Profit & Loss")
    
    # Title Block
    ws_pl.append(["STATEMENT OF PROFIT AND LOSS"])
    ws_pl.cell(1, 1).font = font_title
    
    ws_pl.append([f"For the year ended 31st March, {cy_year}"])
    ws_pl.cell(2, 1).font = font_subtitle
    
    # Blue underline below title
    for c in range(1, 5):
        ws_pl.cell(2, c).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))

    # Unit note (right aligned)
    ws_pl.append([])
    ws_pl.append(["", "", "", f"(All amounts in {profile['currency']} {unit}, unless otherwise stated)"])
    ws_pl.cell(4, 4).font = font_unit_header
    ws_pl.cell(4, 4).alignment = Alignment(horizontal="right")

    # Table Header Rows 5 & 6
    ws_pl.append(["", "Note", f"Year ended", f"Year ended"])
    ws_pl.append(["", "", f"31st March, {cy_year}", f"31st March, {py_year}"])
    
    for r in (5, 6):
        ws_pl.cell(r, 2).font = font_tbl_hdr_note
        ws_pl.cell(r, 2).alignment = Alignment(horizontal="center", vertical="center")
        ws_pl.cell(r, 3).font = font_tbl_hdr_cy
        ws_pl.cell(r, 3).fill = fill_cy_col
        ws_pl.cell(r, 3).alignment = Alignment(horizontal="right", vertical="center")
        ws_pl.cell(r, 4).font = font_tbl_hdr_py
        ws_pl.cell(r, 4).alignment = Alignment(horizontal="right", vertical="center")

    for c in range(1, 5):
        ws_pl.cell(5, c).border = Border(top=Side(style='medium', color=COLOR_BORDER_THIN))
        ws_pl.cell(6, c).border = Border(bottom=Side(style='thin', color=COLOR_BORDER_THIN))

    def _add_pl_line(ws, label, note, cy_val, py_val, is_sec=False, is_subtot=False, is_tot=False, indent="      "):
        lbl_str = f"{indent}{label.strip()}" if (not is_sec and not is_subtot and not is_tot) else label.strip()
        ws.append([lbl_str, str(note) if note else "", cy_val if cy_val != "" else "", py_val if py_val != "" else ""])
        r = ws.max_row
        
        if is_sec:
            ws.cell(r, 1).font = font_major_sec
        elif is_subtot:
            ws.cell(r, 1).font = font_total_lbl
            for c in range(1, 5):
                ws.cell(r, c).border = border_subtotal
        elif is_tot:
            ws.cell(r, 1).font = font_total_lbl
            for c in range(1, 5):
                ws.cell(r, c).border = border_total_dbl
        else:
            ws.cell(r, 1).font = font_item_reg

        ws.cell(r, 2).font = font_tbl_hdr_note
        ws.cell(r, 2).alignment = Alignment(horizontal="center")
        
        ws.cell(r, 3).fill = fill_cy_col
        if isinstance(cy_val, (int, float)):
            ws.cell(r, 3).font = font_total_cy if (is_subtot or is_tot) else font_cy_data
            ws.cell(r, 3).number_format = NUM_FMT
            ws.cell(r, 3).alignment = Alignment(horizontal="right")
            
        if isinstance(py_val, (int, float)):
            ws.cell(r, 4).font = font_total_py if (is_subtot or is_tot) else font_py_data
            ws.cell(r, 4).number_format = NUM_FMT
            ws.cell(r, 4).alignment = Alignment(horizontal="right")

    _add_pl_line(ws_pl, "REVENUE FROM OPERATIONS", "22", data['pl_revenue_ops_cy'], data['pl_revenue_ops_py'], is_sec=True)
    _add_pl_line(ws_pl, "Other income", "23", data['pl_other_income_cy'], data['pl_other_income_py'], indent="    ")
    _add_pl_line(ws_pl, "TOTAL REVENUE", "", data['total_revenue_cy'], data['total_revenue_py'], is_subtot=True)
    
    ws_pl.append([])
    _add_pl_line(ws_pl, "EXPENSES", "", "", "", is_sec=True)
    for exp in data['exp_items']:
        _add_pl_line(ws_pl, exp['title'], exp['note'], exp['cy'], exp['py'], indent="    ")
    _add_pl_line(ws_pl, "TOTAL EXPENSES", "", data['total_expenses_cy'], data['total_expenses_py'], is_subtot=True)

    ws_pl.append([])
    _add_pl_line(ws_pl, "Profit before exceptional items and tax", "", data['pbt_cy'], data['pbt_py'], is_sec=True)
    _add_pl_line(ws_pl, "Exceptional items", "", 0.0, 0.0, indent="    ")
    _add_pl_line(ws_pl, "Profit before tax", "", data['pbt_cy'], data['pbt_py'], is_subtot=True)

    ws_pl.append([])
    _add_pl_line(ws_pl, "Tax expenses", "", "", "", is_sec=True)
    _add_pl_line(ws_pl, "Current tax", "15", -abs(data['tax_curr_cy']), -abs(data['tax_curr_py']), indent="    ")
    _add_pl_line(ws_pl, "Deferred tax credit/(charge)", "15", data['tax_def_cy'], data['tax_def_py'], indent="    ")
    _add_pl_line(ws_pl, "PROFIT FOR THE YEAR", "", data['pat_cy'], data['pat_py'], is_tot=True)

    ws_pl.append([])
    _add_pl_line(ws_pl, "Earnings per equity share", "30", "", "", is_sec=True)
    ws_pl.append(["      Basic (Face value of Rs. 10 each)", "", f"Rs. {data['eps_cy']:,.2f}", f"Rs. {data['eps_py']:,.2f}"])
    r_eps1 = ws_pl.max_row
    ws_pl.cell(r_eps1, 1).font = font_item_reg
    ws_pl.cell(r_eps1, 3).font = font_cy_data
    ws_pl.cell(r_eps1, 3).fill = fill_cy_col
    ws_pl.cell(r_eps1, 3).alignment = Alignment(horizontal="right")
    ws_pl.cell(r_eps1, 4).font = font_py_data
    ws_pl.cell(r_eps1, 4).alignment = Alignment(horizontal="right")

    ws_pl.append(["      Diluted (Face value of Rs. 10 each)", "", f"Rs. {data['eps_cy']:,.2f}", f"Rs. {data['eps_py']:,.2f}"])
    r_eps2 = ws_pl.max_row
    ws_pl.cell(r_eps2, 1).font = font_item_reg
    ws_pl.cell(r_eps2, 3).font = font_cy_data
    ws_pl.cell(r_eps2, 3).fill = fill_cy_col
    ws_pl.cell(r_eps2, 3).alignment = Alignment(horizontal="right")
    ws_pl.cell(r_eps2, 4).font = font_py_data
    ws_pl.cell(r_eps2, 4).alignment = Alignment(horizontal="right")

    _add_statutory_signatures(ws_pl, max_cols=4)

    # ═══════════════════════════════════════════════════════════════════════
    # 3. CASH FLOW STATEMENT (Matches Sample Image 3)
    # ═══════════════════════════════════════════════════════════════════════
    ws_cf = wb.create_sheet(title="Cash Flow Statement")
    
    # Title Block
    ws_cf.append(["CASH FLOW STATEMENT"])
    ws_cf.cell(1, 1).font = font_title
    
    ws_cf.append([f"For the year ended 31st March, {cy_year}"])
    ws_cf.cell(2, 1).font = font_subtitle
    
    for c in range(1, 5):
        ws_cf.cell(2, c).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))

    ws_cf.append([])
    ws_cf.append(["", "", "", f"(All amounts in {profile['currency']} {unit}, unless otherwise stated)"])
    ws_cf.cell(4, 4).font = font_unit_header
    ws_cf.cell(4, 4).alignment = Alignment(horizontal="right")

    # Table Header Rows 5 & 6
    ws_cf.append(["", "Particulars", f"Year ended", f"Year ended"])
    ws_cf.append(["", "", f"31st March, {cy_year}", f"31st March, {py_year}"])
    
    for r in (5, 6):
        ws_cf.cell(r, 2).font = font_tbl_hdr_note
        ws_cf.cell(r, 2).alignment = Alignment(horizontal="left", vertical="center")
        ws_cf.cell(r, 3).font = font_tbl_hdr_cy
        ws_cf.cell(r, 3).fill = fill_cy_col
        ws_cf.cell(r, 3).alignment = Alignment(horizontal="right", vertical="center")
        ws_cf.cell(r, 4).font = font_tbl_hdr_py
        ws_cf.cell(r, 4).alignment = Alignment(horizontal="right", vertical="center")

    for c in range(1, 5):
        ws_cf.cell(5, c).border = Border(top=Side(style='medium', color=COLOR_BORDER_THIN))
        ws_cf.cell(6, c).border = Border(bottom=Side(style='thin', color=COLOR_BORDER_THIN))

    cf = data['cash_flow']
    def _add_cfs_row(ws, sec_code, label, cy_val, py_val="", is_sec=False, is_sub=False, is_subtot=False, is_tot=False, indent="    "):
        ws.append([sec_code, f"{indent}{label}" if (not is_sec and not is_sub and not is_tot) else label, cy_val if cy_val != "" else "", py_val])
        r = ws.max_row
        
        ws.cell(r, 1).font = font_major_sec
        ws.cell(r, 1).alignment = Alignment(horizontal="center")
        
        if is_sec:
            ws.cell(r, 2).font = font_major_sec
        elif is_sub:
            ws.cell(r, 2).font = font_item_sub
        elif is_subtot:
            ws.cell(r, 2).font = font_total_lbl
            for c in range(1, 5):
                ws.cell(r, c).border = border_subtotal
        elif is_tot:
            ws.cell(r, 2).font = font_total_lbl
            for c in range(1, 5):
                ws.cell(r, c).border = border_total_dbl
        else:
            ws.cell(r, 2).font = font_item_reg

        ws.cell(r, 3).fill = fill_cy_col
        if isinstance(cy_val, (int, float)):
            ws.cell(r, 3).font = font_total_cy if (is_subtot or is_tot) else font_cy_data
            ws.cell(r, 3).number_format = NUM_FMT
            ws.cell(r, 3).alignment = Alignment(horizontal="right")
            
        if isinstance(py_val, (int, float)):
            ws.cell(r, 4).font = font_total_py if (is_subtot or is_tot) else font_py_data
            ws.cell(r, 4).number_format = NUM_FMT
            ws.cell(r, 4).alignment = Alignment(horizontal="right")

    _add_cfs_row(ws_cf, "A", "CASH FLOW FROM OPERATING ACTIVITIES:", "", is_sec=True)
    _add_cfs_row(ws_cf, "", "Profit before exceptional items and tax", cf.get('pbt', 0.0), indent="  ")
    _add_cfs_row(ws_cf, "", "Adjustments for:", "", is_sub=True)
    _add_cfs_row(ws_cf, "", "Depreciation and amortisation expenses", cf.get('depreciation', 0.0))
    _add_cfs_row(ws_cf, "", "Finance costs", cf.get('finance_costs', 0.0))
    if cf.get('interest_income', 0) > 0:
        _add_cfs_row(ws_cf, "", "Interest income", -cf['interest_income'])
    
    _add_cfs_row(ws_cf, "", "Cash Generated from operations before working capital changes", cf.get('op_cash_before_wc', 0.0), is_subtot=True, indent="")
    _add_cfs_row(ws_cf, "", "Adjustments for:", "", is_sub=True)
    _add_cfs_row(ws_cf, "", "(Increase)/decrease in inventories", cf.get('inv_change', 0.0))
    _add_cfs_row(ws_cf, "", "(Increase)/decrease in trade receivables", cf.get('rec_change', 0.0))
    _add_cfs_row(ws_cf, "", "(Increase)/decrease in other current assets & loans", round((cf.get('other_ca_change', 0) + cf.get('st_loans_change', 0)), data['decimals']))
    _add_cfs_row(ws_cf, "", "Increase/(decrease) in trade payables", cf.get('pay_change', 0.0))
    _add_cfs_row(ws_cf, "", "Increase/(decrease) in other current liabilities & provisions", cf.get('other_liab_change', 0.0))
    
    _add_cfs_row(ws_cf, "", "Cash generated from operations", cf.get('op_cash_before_wc', 0.0) + cf.get('inv_change', 0.0) + cf.get('rec_change', 0.0) + cf.get('pay_change', 0.0) + cf.get('other_liab_change', 0.0), is_subtot=True, indent="")
    _add_cfs_row(ws_cf, "", "Taxes paid (net of refunds)", -(cf.get('tax_paid', 0.0)), indent="  ")
    _add_cfs_row(ws_cf, "", "Net cash generated from operating activities - [A]", cf.get('net_operating', 0.0), is_subtot=True, indent="")

    ws_cf.append([])
    _add_cfs_row(ws_cf, "B", "CASH FLOW FROM INVESTING ACTIVITIES:", "", is_sec=True)
    _add_cfs_row(ws_cf, "", "Purchase of tangible/intangible assets", cf.get('ppe_purchase', 0.0))
    if cf.get('interest_income', 0) > 0:
        _add_cfs_row(ws_cf, "", "Interest received", cf['interest_income'])
    _add_cfs_row(ws_cf, "", "Net cash used in investing activities - [B]", cf.get('net_investing', 0.0), is_subtot=True, indent="")

    ws_cf.append([])
    _add_cfs_row(ws_cf, "C", "CASH FLOW FROM FINANCING ACTIVITIES:", "", is_sec=True)
    if cf.get('share_capital_proceeds', 0) != 0:
        _add_cfs_row(ws_cf, "", "Proceeds from Issue of Share Capital", cf.get('share_capital_proceeds', 0))
    _add_cfs_row(ws_cf, "", "Proceeds from / (Repayment of) Borrowings", cf.get('borrowings_change', 0.0))
    _add_cfs_row(ws_cf, "", "Finance costs paid", -(cf.get('finance_costs', 0.0)))
    _add_cfs_row(ws_cf, "", "Net cash from / (used in) financing activities - [C]", cf.get('net_financing', 0.0), is_subtot=True, indent="")

    ws_cf.append([])
    _add_cfs_row(ws_cf, "", "Net Increase / (Decrease) in Cash and Cash Equivalents [A + B + C]", cf.get('net_cash_increase', 0.0), is_subtot=True, indent="")
    _add_cfs_row(ws_cf, "", "Cash and cash equivalents at the beginning of the year", cf.get('opening_cash', 0.0), indent="  ")
    _add_cfs_row(ws_cf, "", "Cash and cash equivalents at the end of the year", cf.get('closing_cash', 0.0), is_tot=True, indent="")
    _add_cfs_row(ws_cf, "", "Closing Cash and bank balances per Balance Sheet (Note 19)", cf.get('actual_closing_cash', cf.get('closing_cash', 0.0)), indent="  ")

    # ═══════════════════════════════════════════════════════════════════════
    # 4. ACCOUNTING POLICIES SHEET (Note 1 - Separate Pure Text Sheet)
    # ═══════════════════════════════════════════════════════════════════════
    ws_pol = wb.create_sheet(title="Accounting Policies")
    ws_pol.append(["NOTES"])
    ws_pol.cell(1, 1).font = font_title
    
    ws_pol.append([f"to the financial statements for the year ended 31st March, {cy_year} — Significant Accounting Policies"])
    ws_pol.cell(2, 1).font = font_subtitle
    
    ws_pol.cell(2, 1).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))

    ws_pol.append([])
    ws_pol.append([f"1)  CORPORATE INFORMATION & SIGNIFICANT ACCOUNTING POLICIES"])
    ws_pol.cell(4, 1).font = font_major_sec

    note1_obj = next((n for n in notes_data.get('notes', []) if str(n.get('note_no')) == '1'), None)
    if note1_obj:
        for sec in note1_obj.get('sections', []):
            if sec.get('heading'):
                ws_pol.append([f"    {sec['heading']}"])
                ws_pol.cell(ws_pol.max_row, 1).font = font_sub_head
            
            if sec.get('content'):
                ws_pol.append([f"    {sec['content']}"])
                r_txt = ws_pol.max_row
                ws_pol.cell(r_txt, 1).font = font_item_reg
                ws_pol.cell(r_txt, 1).alignment = Alignment(wrap_text=True)
                ws_pol.append([]) # spacing between clauses

    # ═══════════════════════════════════════════════════════════════════════
    # 5. NOTES TO ACCOUNTS SHEET (Notes 2 to 29 - Matches Sample Image 4)
    # ═══════════════════════════════════════════════════════════════════════
    ws_notes = wb.create_sheet(title="Notes to Accounts")
    
    # Title Block
    ws_notes.append(["NOTES"])
    ws_notes.cell(1, 1).font = font_title
    
    ws_notes.append([f"to the financial statements for the year ended 31st March, {cy_year} (Contd.)"])
    ws_notes.cell(2, 1).font = font_subtitle
    
    for c in range(1, 4):
        ws_notes.cell(2, c).border = Border(bottom=Side(style='medium', color=COLOR_PRIMARY_BLUE))

    ws_notes.append([])
    ws_notes.append(["", "", f"(All amounts in {profile['currency']} {unit}, unless otherwise stated)"])
    ws_notes.cell(4, 3).font = font_unit_header
    ws_notes.cell(4, 3).alignment = Alignment(horizontal="right")

    # Notes 2 to 29 exclusively
    schedule_notes = [n for n in notes_data.get('notes', []) if str(n.get('note_no')) != '1']

    for note in schedule_notes:
        is_approved = (note.get('review_status') == 'Approved') or (not note.get('review_flag'))
        note_banner_text = f"{note.get('note_no')})  {note.get('title', '').upper()}"
        if not is_approved and note.get('review_flag'):
            note_banner_text += "  [⚠️ CA Review Required]"
        
        ws_notes.append([])
        ws_notes.append([note_banner_text, "", ""])
        r_b = ws_notes.max_row
        ws_notes.cell(r_b, 1).font = font_major_sec
        ws_notes.cell(r_b, 1).border = Border(bottom=Side(style='thin', color=COLOR_PRIMARY_BLUE))

        for sec in note.get('sections', []):
            if sec.get('heading'):
                ws_notes.append([f"    {sec['heading']}", "", ""])
                ws_notes.cell(ws_notes.max_row, 1).font = font_sub_head
            
            # Only print internal CA review flag if the note has NOT been approved yet
            if not is_approved and sec.get('ca_review'):
                ws_notes.append([f"    ⚠️ CA Review: {sec['ca_review']}", "", ""])
                ws_notes.cell(ws_notes.max_row, 1).font = font_item_sub

            if sec.get('type') == 'text':
                ws_notes.append([f"    {sec.get('content', '')}", "", ""])
                r_t = ws_notes.max_row
                ws_notes.cell(r_t, 1).font = font_item_reg
                ws_notes.cell(r_t, 1).alignment = Alignment(wrap_text=True)

            elif sec.get('type') == 'table':
                cols = sec.get('columns', [])
                rows = sec.get('rows', [])
                
                # Sub-table headers matching Image 4
                ws_notes.append(["", f"As at", f"As at"])
                ws_notes.append(["", f"31st March, {cy_year}", f"31st March, {py_year}"])
                
                r_h1 = ws_notes.max_row - 1
                r_h2 = ws_notes.max_row
                
                for r in (r_h1, r_h2):
                    ws_notes.cell(r, 2).font = font_tbl_hdr_cy
                    ws_notes.cell(r, 2).fill = fill_cy_col
                    ws_notes.cell(r, 2).alignment = Alignment(horizontal="right", vertical="center")
                    ws_notes.cell(r, 3).font = font_tbl_hdr_py
                    ws_notes.cell(r, 3).alignment = Alignment(horizontal="right", vertical="center")

                for c in range(1, 4):
                    ws_notes.cell(r_h1, c).border = Border(top=Side(style='medium', color=COLOR_BORDER_THIN))
                    ws_notes.cell(r_h2, c).border = Border(bottom=Side(style='thin', color=COLOR_BORDER_THIN))

                for idx, row_data in enumerate(rows):
                    is_tot = sec.get('total_row', False) and (idx == len(rows) - 1)
                    parsed_cy = None
                    parsed_py = None
                    
                    item_lbl = str(row_data[0]) if len(row_data) > 0 else ""
                    
                    if len(row_data) > 1 and str(row_data[1]).strip() not in ('', '—', '-'):
                        try: parsed_cy = float(str(row_data[1]).replace(',', '').strip())
                        except ValueError: parsed_cy = row_data[1]
                        
                    if len(row_data) > 2 and str(row_data[2]).strip() not in ('', '—', '-'):
                        try: parsed_py = float(str(row_data[2]).replace(',', '').strip())
                        except ValueError: parsed_py = row_data[2]

                    ws_notes.append([f"    {item_lbl}", parsed_cy if parsed_cy is not None else "", parsed_py if parsed_py is not None else ""])
                    r_d = ws_notes.max_row
                    
                    if is_tot:
                        ws_notes.cell(r_d, 1).font = font_total_lbl
                        ws_notes.cell(r_d, 2).font = font_total_cy
                        ws_notes.cell(r_d, 2).fill = fill_cy_col
                        ws_notes.cell(r_d, 3).font = font_total_py
                        for c in range(1, 4):
                            ws_notes.cell(r_d, c).border = border_total_dbl
                    else:
                        ws_notes.cell(r_d, 1).font = font_item_reg
                        ws_notes.cell(r_d, 2).font = font_cy_data
                        ws_notes.cell(r_d, 2).fill = fill_cy_col
                        ws_notes.cell(r_d, 3).font = font_py_data
                    
                    if isinstance(parsed_cy, (int, float)):
                        ws_notes.cell(r_d, 2).number_format = NUM_FMT
                        ws_notes.cell(r_d, 2).alignment = Alignment(horizontal="right")
                    if isinstance(parsed_py, (int, float)):
                        ws_notes.cell(r_d, 3).number_format = NUM_FMT
                        ws_notes.cell(r_d, 3).alignment = Alignment(horizontal="right")

        if note.get('additional_remarks'):
            ws_notes.append([])
            ws_notes.append([f"    📝 Additional Auditor Notes / Footnotes (Note {note.get('note_no', '')}):", "", ""])
            r_rem_h = ws_notes.max_row
            ws_notes.cell(r_rem_h, 1).font = font_sub_head
            ws_notes.append([f"    {note['additional_remarks']}", "", ""])
            r_rem_b = ws_notes.max_row
            ws_notes.cell(r_rem_b, 1).font = font_item_sub
            ws_notes.cell(r_rem_b, 1).alignment = Alignment(wrap_text=True)

    # ═══════════════════════════════════════════════════════════════════════
    # 6. TRIAL BALANCE SHEET (Raw Imported Ledgers)
    # ═══════════════════════════════════════════════════════════════════════
    ws_tb = wb.create_sheet(title="Trial Balance")
    ws_tb.append(["TRIAL BALANCE"])
    ws_tb.cell(1, 1).font = font_title
    ws_tb.append([f"Raw Imported Trial Balance Ledgers | FY {profile['financial_year']}"])
    ws_tb.cell(2, 1).font = font_subtitle
    ws_tb.append([])
    
    tb_headers = [
        "Ledger Name", 
        "Tally Parent Group", 
        "Opening Dr (Rs.)", 
        "Opening Cr (Rs.)", 
        "Debit Tx (Rs.)", 
        "Credit Tx (Rs.)", 
        "Closing Dr (Rs.)", 
        "Closing Cr (Rs.)", 
        "Prior Year Net (Rs.)"
    ]
    ws_tb.append(tb_headers)
    for c_idx in range(1, len(tb_headers) + 1):
        ws_tb.cell(4, c_idx).font = font_tbl_hdr_py
        ws_tb.cell(4, c_idx).border = border_tbl_hdr
        ws_tb.cell(4, c_idx).alignment = Alignment(horizontal="center" if c_idx <= 2 else "right")

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT ledger_name, tally_group, opening_dr, opening_cr, debit, credit, closing_dr, closing_cr, prior_closing_net
    FROM tally_import
    ORDER BY tally_group ASC, ledger_name ASC;
    """)
    tb_rows = cursor.fetchall()
    conn.close()

    tot_op_dr = 0.0; tot_op_cr = 0.0; tot_tx_dr = 0.0; tot_tx_cr = 0.0
    tot_cl_dr = 0.0; tot_cl_cr = 0.0; tot_py_net = 0.0

    for r in tb_rows:
        op_dr = r['opening_dr'] or 0.0; op_cr = r['opening_cr'] or 0.0
        tx_dr = r['debit'] or 0.0; tx_cr = r['credit'] or 0.0
        cl_dr = r['closing_dr'] or 0.0; cl_cr = r['closing_cr'] or 0.0
        py_net = r['prior_closing_net'] or 0.0

        tot_op_dr += op_dr; tot_op_cr += op_cr; tot_tx_dr += tx_dr; tot_tx_cr += tx_cr
        tot_cl_dr += cl_dr; tot_cl_cr += cl_cr; tot_py_net += py_net

        ws_tb.append([r['ledger_name'], r['tally_group'], op_dr, op_cr, tx_dr, tx_cr, cl_dr, cl_cr, py_net])
        curr_r = ws_tb.max_row
        ws_tb.cell(curr_r, 1).font = font_item_reg
        ws_tb.cell(curr_r, 2).font = font_item_reg
        for c_idx in range(3, 10):
            ws_tb.cell(curr_r, c_idx).number_format = NUM_FMT
            ws_tb.cell(curr_r, c_idx).font = font_py_data
            ws_tb.cell(curr_r, c_idx).alignment = Alignment(horizontal="right")

    ws_tb.append(["TOTAL TRIAL BALANCE", f"({len(tb_rows)} Ledgers)", tot_op_dr, tot_op_cr, tot_tx_dr, tot_tx_cr, tot_cl_dr, tot_cl_cr, tot_py_net])
    tot_r = ws_tb.max_row
    ws_tb.cell(tot_r, 1).font = font_total_lbl
    ws_tb.cell(tot_r, 2).font = font_total_lbl
    for c_idx in range(1, 10):
        cell = ws_tb.cell(tot_r, c_idx)
        cell.border = border_total_dbl
        if c_idx >= 3:
            cell.font = font_total_py
            cell.number_format = NUM_FMT
            cell.alignment = Alignment(horizontal="right")

    # ═══════════════════════════════════════════════════════════════════════
    # 7. FINANCIAL RATIOS SHEET
    # ═══════════════════════════════════════════════════════════════════════
    ws_rat = wb.create_sheet(title="Financial Ratios")
    ws_rat.append(["MANDATORY SCHEDULE III RATIOS"])
    ws_rat.cell(1, 1).font = font_title
    ws_rat.append([f"11 Statutory Financial Ratios Analysis — FY {profile['financial_year']}"])
    ws_rat.cell(2, 1).font = font_subtitle
    ws_rat.append([])
    
    rat_headers = ["Ratio Name", "Numerator", "Denominator", f"Current Year ({profile['financial_year']})", f"Prior Year ({profile['comparative_year']})", "Variance %", "Benchmark"]
    ws_rat.append(rat_headers)
    for c_idx in range(1, len(rat_headers) + 1):
        ws_rat.cell(4, c_idx).font = font_tbl_hdr_py
        ws_rat.cell(4, c_idx).border = border_tbl_hdr
        ws_rat.cell(4, c_idx).alignment = Alignment(horizontal="center" if c_idx in (4,5,6,7) else "left")

    for r in data['ratios']:
        ws_rat.append([r.get('ratio', r.get('name', '')), r.get('numerator', ''), r.get('denominator', ''), r.get('cy'), r.get('py'), r.get('variance', ''), r.get('benchmark', '')])
        r_idx = ws_rat.max_row
        ws_rat.cell(r_idx, 1).font = font_total_lbl
        ws_rat.cell(r_idx, 2).font = font_item_reg
        ws_rat.cell(r_idx, 3).font = font_item_reg
        if r.get('cy') is not None and isinstance(r.get('cy'), (int, float)):
            ws_rat.cell(r_idx, 4).font = font_cy_data
            ws_rat.cell(r_idx, 4).fill = fill_cy_col
            ws_rat.cell(r_idx, 4).number_format = NUM_FMT
        if r.get('py') is not None and isinstance(r.get('py'), (int, float)):
            ws_rat.cell(r_idx, 5).font = font_py_data
            ws_rat.cell(r_idx, 5).number_format = NUM_FMT
        ws_rat.cell(r_idx, 6).alignment = Alignment(horizontal="right")

    # ═══════════════════════════════════════════════════════════════════════
    # 8. LEDGER MAPPING REGISTER SHEET
    # ═══════════════════════════════════════════════════════════════════════
    ws_map = wb.create_sheet(title="Mapping Register")
    ws_map.append(["LEDGER MAPPING REGISTER"])
    ws_map.cell(1, 1).font = font_title
    ws_map.append([f"Statutory Classification & Note Mapping — FY {profile['financial_year']}"])
    ws_map.cell(2, 1).font = font_subtitle
    ws_map.append([])
    
    map_headers = ["Ledger Name", "Tally Group", f"Closing Balance Net ({unit})", "Schedule III Line Item", "Note No.", "Normal Bal", "Classification", "Cash Flow Category", "CA Review Status"]
    ws_map.append(map_headers)
    for c_idx in range(1, len(map_headers) + 1):
        ws_map.cell(4, c_idx).font = font_tbl_hdr_py
        ws_map.cell(4, c_idx).border = border_tbl_hdr
        ws_map.cell(4, c_idx).alignment = Alignment(horizontal="center" if c_idx in (3,5,6,7,8,9) else "left")

    mapping_rows = get_mapping_register()
    for m in mapping_rows:
        ws_map.append([
            m.get('ledger_name', ''), m.get('tally_group', ''), m.get('closing_net', 0.0),
            m.get('schedule3_head', ''), m.get('note_no', ''), m.get('normal_balance', ''),
            m.get('classification', ''), m.get('cash_flow_category', ''), m.get('review_status', '')
        ])
        r_idx = ws_map.max_row
        ws_map.cell(r_idx, 1).font = font_total_lbl
        ws_map.cell(r_idx, 2).font = font_item_reg
        if isinstance(m.get('closing_net'), (int, float)):
            ws_map.cell(r_idx, 3).font = font_cy_data
            ws_map.cell(r_idx, 3).fill = fill_cy_col
            ws_map.cell(r_idx, 3).number_format = NUM_FMT
        ws_map.cell(r_idx, 5).alignment = Alignment(horizontal="center")
        ws_map.cell(r_idx, 6).alignment = Alignment(horizontal="center")

    # ═══════════════════════════════════════════════════════════════════════
    # Column Widths Optimization
    # ═══════════════════════════════════════════════════════════════════════
    ws_bs.column_dimensions['A'].width = 46
    ws_bs.column_dimensions['B'].width = 8
    ws_bs.column_dimensions['C'].width = 18
    ws_bs.column_dimensions['D'].width = 18

    ws_pl.column_dimensions['A'].width = 46
    ws_pl.column_dimensions['B'].width = 8
    ws_pl.column_dimensions['C'].width = 18
    ws_pl.column_dimensions['D'].width = 18

    ws_cf.column_dimensions['A'].width = 6
    ws_cf.column_dimensions['B'].width = 54
    ws_cf.column_dimensions['C'].width = 18
    ws_cf.column_dimensions['D'].width = 18

    ws_pol.column_dimensions['A'].width = 110

    ws_notes.column_dimensions['A'].width = 50
    ws_notes.column_dimensions['B'].width = 18
    ws_notes.column_dimensions['C'].width = 18

    for sheet in (ws_tb, ws_rat, ws_map):
        for col in sheet.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = get_column_letter(col[0].column)
            sheet.column_dimensions[col_letter].width = min(max(max_len + 3, 13), 50)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    return output_path

def export_word_draft(output_path):
    data = generate_financial_statements()
    notes_data = generate_notes()
    profile = data['profile']

    doc = docx.Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{profile['company_name']}\n")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x4C, 0x97)

    sub = title.add_run(f"STANDALONE FINANCIAL STATEMENTS\nFor the Financial Year {profile['financial_year']}\n")
    sub.font.size = Pt(14)
    sub.font.bold = True

    cin = title.add_run(f"CIN: {profile.get('cin', '—')}\n{profile.get('registered_address', '—')}\n")
    cin.font.size = Pt(10)
    cin.font.italic = True

    doc.add_paragraph("\n")
    doc.add_heading("1. Balance Sheet", level=1)

    # Balance Sheet Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Particulars"
    hdr_cells[1].text = "Note"
    hdr_cells[2].text = f"FY {profile['financial_year']}"
    hdr_cells[3].text = f"FY {profile['comparative_year']}"

    for row in data['bs_equity_liabilities'] + [{'title': 'TOTAL EQUITY AND LIABILITIES', 'is_total': True, 'cy': data['total_eq_liab_cy'], 'py': data['total_eq_liab_py']}] + data['bs_assets'] + [{'title': 'TOTAL ASSETS', 'is_total': True, 'cy': data['total_assets_cy'], 'py': data['total_assets_py']}]:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row.get('title', ''))
        row_cells[1].text = str(row.get('note', ''))
        row_cells[2].text = f"{row.get('cy', 0.0):,.2f}" if 'cy' in row and row.get('cy') is not None else ""
        row_cells[3].text = f"{row.get('py', 0.0):,.2f}" if 'py' in row and row.get('py') is not None else ""

    doc.add_heading("2. Statement of Profit and Loss", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Total Revenue: Rs {data['total_revenue_cy']:,.2f} {data['unit']} (Prior Year: Rs {data['total_revenue_py']:,.2f} {data['unit']})\n").bold = True
    p.add_run(f"Total Expenses: Rs {data['total_expenses_cy']:,.2f} {data['unit']} (Prior Year: Rs {data['total_expenses_py']:,.2f} {data['unit']})\n")
    p.add_run(f"Profit After Tax (PAT): Rs {data['pat_cy']:,.2f} {data['unit']} (Prior Year: Rs {data['pat_py']:,.2f} {data['unit']})\n").bold = True

    # Note 1: Corporate Information & Accounting Policies
    doc.add_heading("3. Corporate Information & Significant Accounting Policies (Note 1)", level=1)
    notes_list = notes_data.get('notes', [])
    note1 = next((n for n in notes_list if n.get('note_no') == '1'), None)
    if note1:
        for sec in note1.get('sections', []):
            if sec.get('heading'):
                h = doc.add_heading(sec['heading'], level=2)
                h.paragraph_format.space_before = Pt(8)
            if sec.get('content'):
                doc.add_paragraph(sec['content'])

    # Notes 2 to 29: Financial Schedules & Additional Auditor Remarks
    doc.add_heading("4. Notes & Financial Schedules (Notes 2 to 29)", level=1)
    for n in [x for x in notes_list if x.get('note_no') != '1']:
        doc.add_heading(f"Note {n.get('note_no')}: {n.get('title')}", level=2)
        for sec in n.get('sections', []):
            if sec.get('heading'):
                doc.add_heading(sec['heading'], level=3)
            if sec.get('type') == 'text':
                doc.add_paragraph(sec.get('content', ''))
            elif sec.get('type') == 'table':
                cols = sec.get('columns', [])
                rows = sec.get('rows', [])
                if cols and rows:
                    t = doc.add_table(rows=1, cols=len(cols))
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, col_name in enumerate(cols):
                        t.rows[0].cells[i].text = str(col_name)
                    for r_data in rows:
                        r_cells = t.add_row().cells
                        for j, val in enumerate(r_data):
                            r_cells[j].text = str(val)
        
        if n.get('additional_remarks'):
            p_rem = doc.add_paragraph()
            r_bold = p_rem.add_run(f"Additional Auditor Remarks & Disclosures (Note {n.get('note_no')}):\n")
            r_bold.bold = True
            r_italic = p_rem.add_run(n['additional_remarks'])
            r_italic.italic = True

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def export_pdf_report(output_path):
    data = generate_financial_statements()
    profile = data['profile']

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#004C97'),
        alignment=1, # Center
        spaceAfter=12
    )

    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        alignment=1,
        spaceAfter=20
    )

    table_cell = ParagraphStyle('TableCell', fontName='Helvetica', fontSize=9, leading=11)
    table_cell_bold = ParagraphStyle('TableCellBold', fontName='Helvetica-Bold', fontSize=9, leading=11)

    story = []

    story.append(Paragraph(profile['company_name'], title_style))
    story.append(Paragraph(f"FINANCIAL STATEMENTS FOR THE YEAR ENDED 31ST MARCH, {profile['financial_year'][-2:]}<br/>(Compliant with Schedule III Division I - Non Ind AS)", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#004C97'), spaceAfter=20))

    # Balance Sheet Table
    story.append(Paragraph("<b>1. BALANCE SHEET AS AT 31ST MARCH, " + profile['financial_year'][-2:] + "</b>", styles['Heading2']))
    story.append(Spacer(1, 8))

    table_data = [
        [Paragraph("<b>Particulars</b>", table_cell_bold), Paragraph("<b>Note</b>", table_cell_bold), Paragraph(f"<b>As at 31.03.{profile['financial_year'][-2:]}</b>", table_cell_bold), Paragraph(f"<b>As at 31.03.{profile['comparative_year'][-2:]}</b>", table_cell_bold)]
    ]

    for row in data['bs_equity_liabilities'] + [{'title': 'TOTAL EQUITY AND LIABILITIES', 'is_total': True, 'cy': data['total_eq_liab_cy'], 'py': data['total_eq_liab_py']}] + data['bs_assets'] + [{'title': 'TOTAL ASSETS', 'is_total': True, 'cy': data['total_assets_cy'], 'py': data['total_assets_py']}]:
        t_style = table_cell_bold if row.get('is_header') or row.get('is_subheader') or row.get('is_total') else table_cell
        table_data.append([
            Paragraph(row['title'], t_style),
            Paragraph(str(row.get('note', '')), table_cell),
            Paragraph(f"{row.get('cy', 0.0):,.2f}" if 'cy' in row else "", t_style),
            Paragraph(f"{row.get('py', 0.0):,.2f}" if 'py' in row else "", t_style)
        ])

    bs_table = Table(table_data, colWidths=[260, 45, 105, 105])
    bs_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#004C97')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E0E0E0')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
    ]))

    story.append(bs_table)
    story.append(PageBreak())

    # Ratios Table
    story.append(Paragraph("<b>2. MANDATORY SCHEDULE III FINANCIAL RATIOS</b>", styles['Heading2']))
    story.append(Spacer(1, 10))

    ratio_data = [
        [Paragraph("<b>Ratio Name</b>", table_cell_bold), Paragraph("<b>Current Year</b>", table_cell_bold), Paragraph("<b>Prior Year</b>", table_cell_bold), Paragraph("<b>Variance %</b>", table_cell_bold)]
    ]
    for r in data['ratios']:
        ratio_data.append([
            Paragraph(r['ratio'], table_cell),
            Paragraph(str(r['cy']), table_cell_bold),
            Paragraph(str(r['py']), table_cell),
            Paragraph(r['variance'], table_cell)
        ])

    r_table = Table(ratio_data, colWidths=[220, 100, 100, 95])
    r_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#004C97')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E0E0E0')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
    ]))
    story.append(r_table)

    doc.build(story)
    return output_path
